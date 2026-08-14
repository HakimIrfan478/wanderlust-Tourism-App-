"""Bundle the project's first-party source into a single submission file.

    python manage.py export_source                  # writes .txt
    python manage.py export_source --format docx    # writes .docx
    python manage.py export_source --format both
    python manage.py export_source --list           # show what would be included

Third-party code is excluded: virtual environments, node_modules, lock files,
build output, caches and binaries. Django-generated migrations are included but
clearly marked as generated, so a reader can tell them apart from hand-written
work rather than having to guess.

The output carries a manifest at the top — file count, line count and the
exclusion rules — so a marker can see at a glance what was and was not included.
"""
from datetime import datetime
from pathlib import Path

from django.core.management.base import BaseCommand, CommandError

PROJECT_ROOT = Path(__file__).resolve().parents[4]

# Directories never descended into.
EXCLUDE_DIRS = {
    ".venv", "venv", "env", "node_modules", "__pycache__", ".git", ".expo",
    ".idea", ".vscode", "dist", "build", "web-build", "staticfiles",
    "evaluation_results", "figures", ".pytest_cache", ".claude",
}

# Files excluded by exact name.
EXCLUDE_FILES = {
    "package-lock.json",   # generated dependency graph
    "yarn.lock",
    "db.sqlite3",
    ".env",                # secrets, and machine-specific
}

# Only these extensions are treated as source.
INCLUDE_SUFFIXES = {".py", ".js", ".jsx", ".json", ".yml", ".yaml", ".example"}

# Files that are data or configuration rather than authored logic. Listed in
# the manifest but not inlined, to keep the submission to actual source.
DATA_FILES = {"destinations.json"}

SECTIONS = [
    ("Backend — project configuration", ["backend/wanderlust"]),
    ("Backend — accounts app", ["backend/accounts"]),
    ("Backend — destinations app", ["backend/destinations"]),
    ("Backend — recommendations app (the two models)", ["backend/recommendations"]),
    ("Backend — evaluation app (the experiment)", ["backend/evaluation"]),
    ("Backend — integrations app (external APIs)", ["backend/integrations"]),
    ("Backend — entry point and dependencies", ["backend/manage.py", "backend/requirements.txt",
                                                "backend/requirements-ci.txt", "backend/.env.example"]),
    ("Frontend — entry points", ["frontend/index.js", "frontend/App.js",
                                 "frontend/app.json", "frontend/package.json",
                                 "frontend/babel.config.js"]),
    ("Frontend — application source", ["frontend/src"]),
    ("Frontend — tooling", ["frontend/scripts"]),
    ("Continuous integration", [".github"]),
]


class Command(BaseCommand):
    help = "Bundle first-party source code into a single submission file."

    def add_arguments(self, parser):
        parser.add_argument("--format", choices=["txt", "docx", "both"], default="txt")
        parser.add_argument("--output-dir", default=None)
        parser.add_argument("--author", default="<your full name>")
        parser.add_argument("--student-number", default="<student number>")
        parser.add_argument("--list", action="store_true", help="List files, write nothing.")

    def handle(self, *args, **options):
        files = self._collect()
        if not files:
            raise CommandError("No source files found — check PROJECT_ROOT.")

        total_lines = sum(len(self._read(p).splitlines()) for _, p in files)

        if options["list"]:
            for section, path in files:
                rel = path.relative_to(PROJECT_ROOT).as_posix()
                self.stdout.write(f"  [{section[:28]:<28}] {rel}")
            self.stdout.write(
                self.style.SUCCESS(f"\n{len(files)} files, {total_lines:,} lines.")
            )
            return

        out_dir = Path(options["output_dir"] or PROJECT_ROOT / "docs" / "submission")
        out_dir.mkdir(parents=True, exist_ok=True)

        written = []
        if options["format"] in ("txt", "both"):
            written.append(self._write_txt(files, total_lines, out_dir, options))
        if options["format"] in ("docx", "both"):
            written.append(self._write_docx(files, total_lines, out_dir, options))

        for path in written:
            size = path.stat().st_size / 1024
            self.stdout.write(self.style.SUCCESS(f"  wrote {path.name}  ({size:,.0f} KB)"))
        self.stdout.write(f"\n{len(files)} files, {total_lines:,} lines, in {out_dir}")

    # -- collection --------------------------------------------------------
    def _collect(self):
        """Return [(section_title, path)] in the order defined by SECTIONS."""
        collected = []
        seen = set()
        for title, targets in SECTIONS:
            for target in targets:
                base = PROJECT_ROOT / target
                if base.is_file():
                    if self._keep(base) and base not in seen:
                        collected.append((title, base))
                        seen.add(base)
                elif base.is_dir():
                    for path in sorted(base.rglob("*")):
                        if path.is_file() and self._keep(path) and path not in seen:
                            collected.append((title, path))
                            seen.add(path)
        return collected

    def _keep(self, path: Path) -> bool:
        if any(part in EXCLUDE_DIRS for part in path.parts):
            return False
        if path.name in EXCLUDE_FILES or path.name in DATA_FILES:
            return False
        if path.suffix not in INCLUDE_SUFFIXES:
            return False
        # Empty package markers carry no information.
        if path.name == "__init__.py" and not path.read_text(encoding="utf-8", errors="replace").strip():
            return False
        return True

    def _read(self, path: Path) -> str:
        return path.read_text(encoding="utf-8", errors="replace")

    def _is_generated(self, path: Path) -> bool:
        return "migrations" in path.parts and path.name != "__init__.py"

    # -- manifest ----------------------------------------------------------
    def _manifest_lines(self, files, total_lines, options):
        by_section = {}
        for section, path in files:
            by_section.setdefault(section, []).append(path)

        lines = [
            "=" * 78,
            "  WANDERLUST — SOURCE CODE SUBMISSION",
            "=" * 78,
            "",
            f"  Author          : {options['author']}",
            f"  Student number  : {options['student_number']}",
            f"  Generated       : {datetime.now().strftime('%Y-%m-%d %H:%M')}",
            f"  Files           : {len(files)}",
            f"  Lines           : {total_lines:,}",
            "",
            "  EXCLUDED FROM THIS SUBMISSION",
            "  Third-party code is not included. Specifically excluded: Python",
            "  virtual environments, node_modules, dependency lock files, build",
            "  output, caches, binary assets, the database file, and environment",
            "  files containing secrets. Dependencies are declared in",
            "  requirements.txt and package.json, both of which are included.",
            "",
            "  Django-generated database migrations ARE included, and each is",
            "  marked as generated so it can be distinguished from hand-written",
            "  code. The destination catalogue (destinations.json) is authored",
            "  data rather than source and is omitted; it is in the repository at",
            "  backend/destinations/data/destinations.json.",
            "",
            "=" * 78,
            "  CONTENTS",
            "=" * 78,
            "",
        ]

        for section in by_section:
            section_lines = sum(len(self._read(p).splitlines()) for p in by_section[section])
            lines.append(f"  {section}  ({len(by_section[section])} files, {section_lines:,} lines)")
            for path in by_section[section]:
                rel = path.relative_to(PROJECT_ROOT).as_posix()
                tag = "  [generated]" if self._is_generated(path) else ""
                lines.append(f"      {rel}{tag}")
            lines.append("")
        return lines

    # -- writers -----------------------------------------------------------
    def _write_txt(self, files, total_lines, out_dir, options):
        parts = self._manifest_lines(files, total_lines, options)

        current = None
        for section, path in files:
            if section != current:
                current = section
                parts += ["", "=" * 78, f"  {section.upper()}", "=" * 78, ""]

            rel = path.relative_to(PROJECT_ROOT).as_posix()
            parts += ["", "-" * 78, f"  FILE: {rel}"]
            if self._is_generated(path):
                parts.append("  NOTE: generated by Django's makemigrations, not hand-written.")
            parts += ["-" * 78, ""]
            parts += self._read(path).splitlines()

        parts += ["", "=" * 78, "  END OF SOURCE CODE SUBMISSION", "=" * 78]

        target = out_dir / "Wanderlust-Source-Code.txt"
        target.write_text("\n".join(parts), encoding="utf-8")
        return target

    def _write_docx(self, files, total_lines, out_dir, options):
        try:
            from docx import Document
            from docx.enum.text import WD_BREAK
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            from docx.shared import Cm, Pt, RGBColor
        except ImportError as exc:
            raise CommandError("python-docx is required for --format docx") from exc

        doc = Document()
        sec = doc.sections[0]
        sec.left_margin = sec.right_margin = Cm(1.6)
        sec.top_margin = sec.bottom_margin = Cm(1.6)

        style = doc.styles["Normal"]
        style.font.name = "Consolas"
        style.font.size = Pt(7.5)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = 1.0

        def mono(text, size=7.5, bold=False, colour=None):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(text if text.strip() else " ")
            run.font.name = "Consolas"
            run.font.size = Pt(size)
            run.bold = bold
            if colour:
                run.font.color.rgb = colour
            rpr = run._element.get_or_add_rPr()
            fonts = OxmlElement("w:rFonts")
            fonts.set(qn("w:ascii"), "Consolas")
            fonts.set(qn("w:hAnsi"), "Consolas")
            rpr.append(fonts)
            return p

        for line in self._manifest_lines(files, total_lines, options):
            mono(line, size=8.5, bold=line.startswith("=") or "WANDERLUST" in line)

        current = None
        for section, path in files:
            if section != current:
                current = section
                doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
                mono("=" * 78, size=9, bold=True)
                mono(f"  {section.upper()}", size=9, bold=True)
                mono("=" * 78, size=9, bold=True)
                mono("")

            rel = path.relative_to(PROJECT_ROOT).as_posix()
            mono("")
            mono("-" * 78, size=8)
            mono(f"  FILE: {rel}", size=8.5, bold=True, colour=RGBColor(0x0A, 0x5C, 0x5B))
            if self._is_generated(path):
                mono("  NOTE: generated by Django's makemigrations, not hand-written.",
                     size=8, colour=RGBColor(0x99, 0x55, 0x00))
            mono("-" * 78, size=8)
            mono("")
            for line in self._read(path).splitlines():
                mono(line)

        target = out_dir / "Wanderlust-Source-Code.docx"
        doc.save(str(target))
        return target
